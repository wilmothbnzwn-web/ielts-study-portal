#!/usr/bin/env python3
"""
IELTS Vocab Extraction & Deck Builder Pipeline
==============================================
Extracts vocabulary from local IELTS resource files (XLSX, DOCX, PDF, XLS),
cleans and structures into unified schema, then chunks into 24-word decks.

Usage:
  python3 scripts/extract_vocab_decks.py [--dry-run] [--max-words 5000]
"""

import json
import os
import re
import sys
import hashlib
import traceback
from datetime import datetime
from collections import OrderedDict

# ── Config ────────────────────────────────────────────────────────
DECK_SIZE = 24          # Target words per deck
MIN_DECK_SIZE = 18      # Minimum words per deck
MAX_DECK_SIZE = 30      # Maximum words per deck

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(PROJECT_ROOT, 'data')
VOCAB_JSON = os.path.join(DATA_DIR, 'vocabulary.json')
OUTPUT_JSON = os.path.join(DATA_DIR, 'vocabulary.json')

IELTS_LIBRARY = os.path.expanduser(
    "~/Desktop/IELTS_Organized_Library/"
    "05_高频词汇与替换语料库_Vocab_and_Synonyms"
)

# ── Source Files ───────────────────────────────────────────────────
SOURCE_FILES = {
    # XLSX files (fast, clean)
    'listen_scene': os.path.join(IELTS_LIBRARY, '01_同义词替换合集', '听力核心场景词汇.xlsx'),
    '179_syn': os.path.join(IELTS_LIBRARY, '03_考点词系列_538_179', '179听力考点词', '听力179词同义替换.xlsx'),
    '538_read': os.path.join(IELTS_LIBRARY, '03_考点词系列_538_179', '538阅读考点词', '雅思阅读-538考点词表格整理.xlsx'),
    # DOCX files
    'taozi_syn': os.path.join(IELTS_LIBRARY, '01_同义词替换合集', '桃子99组雅思阅读同义词替换.docx'),
    'speaking': os.path.join(IELTS_LIBRARY, '01_同义词替换合集', '雅思口语练习题库.docx'),
    'map_vocab': os.path.join(IELTS_LIBRARY, '01_同义词替换合集', '雅思听力地图题词汇汇总.docx'),
    # XLS files (large, slow)
    'xls_order': os.path.join(IELTS_LIBRARY, '02_词汇书与词表', '雅思英语单词表', '雅思词汇EXCEL版-顺序版.xls'),
    'xls_random': os.path.join(IELTS_LIBRARY, '02_词汇书与词表', '雅思英语单词表', '雅思词汇EXCEL词-乱序版.xls'),
}

POS_MAP = {
    'n': 'noun', 'v': 'verb', 'adj': 'adjective', 'adv': 'adverb',
    'prep': 'preposition', 'conj': 'conjunction', 'pron': 'pronoun',
    'num': 'number', 'art': 'article', 'int': 'interjection',
    'noun': 'noun', 'verb': 'verb', 'adjective': 'adjective', 'adverb': 'adverb',
}

THEME_MAP = {
    'rental': 'Social Development',
    'accommodation': 'Social Development',
    'public facilities': 'Environment & Development',
    'facility': 'Environment & Development',
    'transport': 'Social Development',
    'travel': 'Social Development',
    'tourism': 'Social Development',
    'health': 'Health & Medicine',
    'medicine': 'Health & Medicine',
    'education': 'Education & Knowledge',
    'library': 'Education & Knowledge',
    'job': 'Economy & Trade',
    'employment': 'Economy & Trade',
    'business': 'Economy & Business',
    'banking': 'Economy & Finance',
    'shopping': 'Economy & Trade',
    'food': 'Culture & Society',
    'environment': 'Environment & Development',
    'technology': 'Technology & Society',
    'science': 'Academic & Research',
    'culture': 'Culture & Society',
    'history': 'Culture & Society',
    'art': 'Culture & Society',
    'sport': 'Culture & Society',
    'entertainment': 'Culture & Society',
    'media': 'Technology & Society',
    'reading': 'Academic & Research',
    'listening': 'Academic & Research',
}


def guess_theme(text):
    """Guess theme from text content."""
    text_lower = (text or '').lower()
    for keyword, theme in THEME_MAP.items():
        if keyword in text_lower:
            return theme
    return 'Academic & Research'


def extract_pos(chinese_text):
    """Try to extract POS from Chinese description."""
    if not chinese_text:
        return 'noun'
    text = chinese_text.strip()
    # Common patterns in Chinese vocab: "v.调整" or "adj.基本的"
    m = re.match(r'^\s*(n|v|adj|adv|prep|conj|pron|num|art|int)\.?\s*', text, re.IGNORECASE)
    if m:
        pos_abbr = m.group(1).lower().rstrip('.')
        return POS_MAP.get(pos_abbr, 'noun')
    # Match full POS names
    for full_pos in ['adjective', 'adverb', 'preposition', 'conjunction', 'pronoun', 'interjection']:
        if text.lower().startswith(full_pos):
            return full_pos
    return 'noun'


def clean_chinese(text):
    """Clean Chinese text: remove POS prefix, strip whitespace."""
    if not text:
        return ''
    text = text.strip()
    # Remove POS prefixes like "v.", "adj.", "n."
    text = re.sub(r'^\s*(n|v|adj|adv|prep|conj|pron|num|art|int)\.?\s*', '', text, flags=re.IGNORECASE)
    # Also remove full-pos prefixes
    for pos in ['adjective', 'adverb', 'preposition', 'conjunction', 'pronoun', 'interjection']:
        if text.lower().startswith(pos + '.'):
            text = text[len(pos)+1:].strip()
    # Remove leading punctuation/whitespace
    text = re.sub(r'^[;；，,.\s]+', '', text)
    return text.strip()


def word_hash(word):
    """Generate a stable hash for dedup."""
    return hashlib.md5(word.lower().strip().encode()).hexdigest()


# ═══════════════════════════════════════════════════════════════════
# PHASE 1: Extract from XLSX files
# ═══════════════════════════════════════════════════════════════════

def extract_listen_scene(filepath):
    """Extract from 听力核心场景词汇.xlsx (656 words with scene categories)."""
    print(f"  📊 Extracting: 听力核心场景词汇.xlsx ...")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
        ws = wb.active
        words = []
        scene_counts = {}

        for row in ws.iter_rows(min_row=2, values_only=True):
            seq, scene, word, chinese = row[0], row[1], row[2], row[3]
            if not word or not str(word).strip():
                continue
            word_text = str(word).strip()
            chinese_text = str(chinese).strip() if chinese else ''
            scene_text = str(scene).strip() if scene else ''

            if len(word_text) > 50 or not re.search(r'[a-zA-Z]', word_text):
                continue  # Skip non-word rows

            # Track scene counts
            scene_short = scene_text.split(' ')[0] if scene_text else 'General'
            scene_counts[scene_short] = scene_counts.get(scene_short, 0) + 1

            theme = guess_theme(scene_text)

            words.append({
                'word': word_text,
                'pos': 'noun',
                'chinese': chinese_text,
                'synonyms': [],
                'definition': f'IELTS Listening vocabulary — {scene_text}',
                'example': '',
                'theme': theme,
                'difficulty': 2,
                'source': '听力核心场景词汇.xlsx',
                'source_file': '听力核心场景词汇.xlsx',
                'scene': scene_text,
            })

        wb.close()
        print(f"     ✓ {len(words)} words extracted (scenes: {len(scene_counts)})")
        return words
    except Exception as e:
        print(f"     ✗ ERROR: {e}")
        return []


def extract_179_syn(filepath):
    """Extract from 听力179词同义替换.xlsx (179 synonym pairs)."""
    print(f"  📊 Extracting: 听力179词同义替换.xlsx ...")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
        ws = wb.active
        words = []

        for row in ws.iter_rows(min_row=3, values_only=True):
            seq, word_en, chinese, synonyms_str = row[0], row[1], row[2], row[3]
            if not word_en or not str(word_en).strip():
                continue
            word_text = str(word_en).strip()
            if len(word_text) > 50 or not re.search(r'[a-zA-Z]', word_text):
                continue

            chinese_text = str(chinese).strip() if chinese else ''
            syn_text = str(synonyms_str).strip() if synonyms_str else ''

            # Parse synonyms: comma or semicolon separated
            syn_list = []
            if syn_text:
                # Split on comma, semicolon
                parts = re.split(r'[,;，；]\s*', syn_text)
                syn_list = [s.strip() for s in parts if s.strip() and len(s.strip()) < 50][:6]

            words.append({
                'word': word_text,
                'pos': 'verb' if chinese_text and any(v in chinese_text for v in ['动', '使', '做']) else 'noun',
                'chinese': chinese_text,
                'synonyms': syn_list,
                'definition': f'IELTS Listening high-frequency synonym for "{word_text}"',
                'example': '',
                'theme': 'Academic & Research',
                'difficulty': 3,
                'source': '听力179词同义替换.xlsx',
                'source_file': '听力179词同义替换.xlsx',
            })

        wb.close()
        print(f"     ✓ {len(words)} words extracted")
        return words
    except Exception as e:
        print(f"     ✗ ERROR: {e}")
        return []


def extract_538_read(filepath):
    """Extract from 雅思阅读-538考点词表格整理.xlsx (538 reading test-point words)."""
    print(f"  📊 Extracting: 雅思阅读-538考点词表格整理.xlsx ...")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
        ws = wb.active
        words = []

        for row in ws.iter_rows(min_row=3, values_only=True):
            seq, word_en, chinese_raw, synonyms_str = row[0], row[1], row[2], row[3]
            if not word_en or not str(word_en).strip():
                continue
            word_text = str(word_en).strip()
            if len(word_text) > 50:
                continue

            chinese_raw = str(chinese_raw).strip() if chinese_raw else ''
            pos = extract_pos(chinese_raw)
            chinese_text = clean_chinese(chinese_raw)

            syn_text = str(synonyms_str).strip() if synonyms_str else ''
            syn_list = []
            if syn_text:
                parts = re.split(r'[,;，；]\s*', syn_text)
                syn_list = [s.strip() for s in parts if s.strip() and len(s.strip()) < 50][:6]

            theme = guess_theme(chinese_text) if chinese_text else 'Academic & Research'
            difficulty = 4 if len(syn_list) >= 3 else 3

            words.append({
                'word': word_text,
                'pos': pos,
                'chinese': chinese_text,
                'synonyms': syn_list,
                'definition': f'IELTS Reading test-point word: {word_text}',
                'example': '',
                'theme': theme,
                'difficulty': difficulty,
                                'source': '雅思阅读-538考点词表格整理.xlsx',
                'source_file': '雅思阅读-538考点词表格整理.xlsx',
            })

        wb.close()
        print(f"     ✓ {len(words)} words extracted")
        return words
    except Exception as e:
        print(f"     ✗ ERROR: {e}")
        return []


# ═══════════════════════════════════════════════════════════════════
# PHASE 2: Extract from DOCX files
# ═══════════════════════════════════════════════════════════════════

def extract_docx_table(filepath, source_name):
    """Generic DOCX extractor — handles tables AND paragraph-based vocab lists."""
    print(f"  📄 Extracting: {source_name} ...")
    try:
        from docx import Document
        doc = Document(filepath)
        words = []

        # ── Strategy 1: Tables ──
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    eng_idx = None
                    chn_idx = None
                    for ci, cell_text in enumerate(cells):
                        if re.search(r'[a-zA-Z]', cell_text) and not re.search(r'[一-鿿]', cell_text):
                            eng_idx = ci
                        elif re.search(r'[一-鿿]', cell_text):
                            chn_idx = ci
                    if eng_idx is not None:
                        word_text = cells[eng_idx].strip()
                        if len(word_text) < 2 or len(word_text) > 50:
                            continue
                        chinese_text = cells[chn_idx].strip() if chn_idx is not None else ''
                        syn_list = []
                        for ci, cell_text in enumerate(cells):
                            if ci != eng_idx and ci != chn_idx:
                                if re.search(r'[a-zA-Z]', cell_text) and ',' in cell_text:
                                    syn_list = [s.strip() for s in re.split(r'[,;，；]\s*', cell_text) if s.strip()][:6]
                        words.append({
                            'word': word_text, 'pos': 'noun', 'chinese': chinese_text,
                            'synonyms': syn_list,
                            'definition': f'IELTS vocabulary from {source_name}',
                            'example': '', 'theme': guess_theme(chinese_text) if chinese_text else 'Academic & Research',
                            'difficulty': 3, 'source': source_name, 'source_file': source_name,
                        })

        # ── Strategy 2: Paragraph-based extraction ──
        if not words:
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

            # Pattern A: "number. EnglishWord ChineseMeaning / synonym1 / synonym2"
            # e.g. "1. capacity 能力，容量/competence(竞争)能力/ability 能力(笼统的)"
            pattern_a = re.findall(
                r'(?:^|\n)\s*(\d+)[\.、]\s*([a-zA-Z][a-zA-Z\s\-/]{2,40})\s+([一-鿿][一-鿿\s，,、；;.·/()（）a-zA-Z\-]{2,100}?)(?=\n\s*\d+[\.、]|\n\s*[a-zA-Z]|\Z)',
                full_text
            )
            for match in pattern_a:
                seq, word_text, rest = match
                word_text = word_text.strip()
                if len(word_text) < 2 or len(word_text) > 50:
                    continue

                # Split Chinese meaning from synonyms (synonyms are English words)
                chinese_parts = []
                syn_parts = []
                # Split by / or ,
                segments = re.split(r'\s*[/／]\s*|\s*[,，]\s*', rest)
                for seg in segments:
                    seg = seg.strip()
                    if re.search(r'[一-鿿]', seg):
                        chinese_parts.append(seg)
                    elif re.search(r'[a-zA-Z]', seg) and len(seg) > 1:
                        syn_parts.append(seg)

                chinese_text = '；'.join(chinese_parts) if chinese_parts else rest.strip()
                syn_list = syn_parts[:6]

                words.append({
                    'word': word_text, 'pos': 'noun', 'chinese': chinese_text,
                    'synonyms': syn_list,
                    'definition': f'IELTS Reading synonym from {source_name}',
                    'example': '', 'theme': 'Academic & Research', 'difficulty': 4,
                    'source': source_name, 'source_file': source_name,
                })

            # Pattern B: English phrase — Chinese (for map/location vocab)
            # e.g. "on the left/ right 在左边/右边"
            if not words:
                pattern_b = re.findall(
                    r'(?:^|\n)\s*(\d+)[\.、]?\s*([a-zA-Z][a-zA-Z\s\-/]{3,60})\s+([一-鿿][一-鿿\s，,、；;./()（）a-zA-Z\-]{2,80}?)(?=\n\s*\d+[\.、]|\Z)',
                    full_text
                )
                for match in pattern_b:
                    seq, word_text, chinese_text = match
                    word_text = word_text.strip()
                    chinese_text = chinese_text.strip()
                    if len(word_text) < 2 or len(word_text) > 60:
                        continue
                    words.append({
                        'word': word_text, 'pos': 'phrase', 'chinese': chinese_text,
                        'synonyms': [],
                        'definition': f'IELTS vocabulary from {source_name}',
                        'example': '', 'theme': guess_theme(chinese_text) if chinese_text else 'Academic & Research',
                        'difficulty': 2, 'source': source_name, 'source_file': source_name,
                    })

        print(f"     ✓ {len(words)} words extracted")
        return words
    except Exception as e:
        print(f"     ✗ ERROR: {e}")
        traceback.print_exc()
        return []


# ═══════════════════════════════════════════════════════════════════
# PHASE 3: Extract from PDF files
# ═══════════════════════════════════════════════════════════════════

def extract_pdf_text(filepath, source_name, max_pages=30):
    """Extract text from a PDF using pdfplumber."""
    print(f"  📑 Extracting: {source_name} ...")
    try:
        import pdfplumber
        words = []

        with pdfplumber.open(filepath) as pdf:
            pages = pdf.pages[:max_pages]
            for page in pages:
                text = page.extract_text()
                if not text:
                    continue

                # Try to find word-definition patterns in the text
                lines = text.strip().split('\n')
                for line in lines:
                    line = line.strip()
                    if not line or len(line) > 200:
                        continue

                    # Pattern: English word followed by Chinese
                    # e.g., "abandon 放弃；遗弃"
                    m = re.match(r'^([a-zA-Z][a-zA-Z\s\-/]{2,40})\s{2,}([一-鿿；;，,、（()）.\s\-]+)', line)
                    if m:
                        word_text = m.group(1).strip()
                        chinese_text = m.group(2).strip()
                        if ' ' in word_text and len(word_text.split()) > 3:
                            continue  # Too many words, likely a sentence
                        words.append({
                            'word': word_text,
                            'pos': 'noun',
                            'chinese': chinese_text,
                            'synonyms': [],
                            'definition': f'IELTS vocabulary from {source_name}',
                            'example': '',
                            'theme': guess_theme(chinese_text),
                            'difficulty': 2,
                                'source': source_name, 'source_file': source_name,
                        })

                    # Also try CSV-like lines: "word,chinese"
                    elif re.match(r'^[a-zA-Z][a-zA-Z\s\-]{2,30},[一-鿿]', line):
                        parts = line.split(',', 1)
                        word_text = parts[0].strip()
                        chinese_text = parts[1].strip() if len(parts) > 1 else ''
                        words.append({
                            'word': word_text,
                            'pos': 'noun',
                            'chinese': chinese_text,
                            'synonyms': [],
                            'definition': f'IELTS vocabulary from {source_name}',
                            'example': '',
                            'theme': guess_theme(chinese_text),
                            'difficulty': 2,
                                'source': source_name, 'source_file': source_name,
                        })

        print(f"     ✓ {len(words)} words extracted from PDF")
        return words
    except Exception as e:
        print(f"     ✗ ERROR: {e}")
        return []


# ═══════════════════════════════════════════════════════════════════
# PHASE 4: Extract from XLS files (large, use xlrd)
# ═══════════════════════════════════════════════════════════════════

def extract_xls_vocab(filepath, source_name, max_rows=5000):
    """Extract from large XLS vocabulary files."""
    print(f"  📗 Extracting: {source_name} (large XLS, may be slow)...")
    try:
        import xlrd
        wb = xlrd.open_workbook(filepath)
        ws = wb.sheet_by_index(0)
        print(f"     Sheet: {ws.nrows} rows, {ws.ncols} cols")

        words = []
        limit = min(ws.nrows, max_rows + 1)

        for i in range(1, limit):
            vals = ws.row_values(i)
            if not vals:
                continue

            # Common XLS vocab formats: col0=word, col1=chinese | col0=index, col1=word, col2=chinese
            word_text = ''
            chinese_text = ''

            for vi, val in enumerate(vals):
                s = str(val).strip() if val else ''
                if re.search(r'[a-zA-Z]', s) and not re.search(r'[一-鿿]', s) and len(s) >= 2 and len(s) <= 50:
                    if not word_text:
                        word_text = s
                elif re.search(r'[一-鿿]', s) and not re.search(r'[a-zA-Z]{3,}', s):
                    if not chinese_text:
                        chinese_text = s

            if word_text and len(word_text) >= 2:
                pos = 'noun'
                if chinese_text:
                    pos = extract_pos(chinese_text)
                    chinese_text = clean_chinese(chinese_text)

                words.append({
                    'word': word_text,
                    'pos': pos,
                    'chinese': chinese_text,
                    'synonyms': [],
                    'definition': f'IELTS vocabulary from {source_name}',
                    'example': '',
                    'theme': guess_theme(chinese_text) if chinese_text else 'Academic & Research',
                    'difficulty': 2,
                        'source': source_name, 'source_file': source_name,
                })

        print(f"     ✓ {len(words)} words extracted")
        return words
    except Exception as e:
        print(f"     ✗ ERROR: {e}")
        traceback.print_exc()
        return []


# ═══════════════════════════════════════════════════════════════════
# PHASE 5: Merge, deduplicate, assign IDs
# ═══════════════════════════════════════════════════════════════════

def normalize_word_entry(entry):
    """Ensure all required fields exist."""
    return {
        'word': str(entry.get('word', '')).strip(),
        'pos': entry.get('pos', 'noun'),
        'chinese': str(entry.get('chinese', '')).strip(),
        'synonyms': entry.get('synonyms', []) or [],
        'definition': str(entry.get('definition', '')).strip(),
        'example': str(entry.get('example', '')).strip(),
        'theme': entry.get('theme', 'Academic & Research'),
        'difficulty': int(entry.get('difficulty', 2)),
        'source': str(entry.get('source', entry.get('source_file', 'unknown'))), 'source_file': str(entry.get('source_file', 'unknown')),
    }


def is_valid_entry(entry):
    """Filter out invalid entries."""
    w = entry.get('word', '')
    if not w or len(w) < 2 or len(w) > 50:
        return False
    # Must contain at least some letters
    if not re.search(r'[a-zA-Z]', w):
        return False
    # Filter out clearly non-word entries
    if re.search(r'^\d+$', w):  # Pure numbers
        return False
    if len(w.split()) > 5:  # Too many words (likely a sentence)
        return False
    # Valid Chinese should exist
    if not entry.get('chinese'):
        return False
    return True


def merge_and_dedup(new_words, existing_words):
    """Merge new words with existing, deduplicate by word."""
    seen = OrderedDict()

    # Load existing first (they have correct IDs)
    for w in existing_words:
        key = w['word'].lower().strip()
        seen[key] = w

    # Add new words (skip if already exists)
    added = 0
    skipped = 0
    for w in new_words:
        key = w['word'].lower().strip()
        if key not in seen:
            seen[key] = w
            added += 1
        else:
            skipped += 1
            # Merge synonyms if existing entry has fewer
            existing = seen[key]
            if w.get('synonyms') and (not existing.get('synonyms') or len(existing.get('synonyms', [])) < len(w.get('synonyms', []))):
                existing['synonyms'] = w['synonyms']
            # Use better definition if available
            if w.get('definition') and 'vocabulary from' not in w.get('definition', '') and ('vocabulary from' in existing.get('definition', '') or not existing.get('definition')):
                existing['definition'] = w['definition']

    # Re-assign IDs
    result = []
    for i, entry in enumerate(seen.values()):
        entry['id'] = i + 1
        result.append(entry)

    return result, added, skipped


# ═══════════════════════════════════════════════════════════════════
# PHASE 6: Chunk into decks (matching js/vocab-decks.js logic)
# ═══════════════════════════════════════════════════════════════════

def slugify(text):
    return re.sub(r'[^a-z0-9]+', '-', str(text).lower()).strip('-')


def build_decks(words):
    """Group words by source_file and chunk into 24-word decks."""

    # Group by source_file
    source_groups = OrderedDict()
    for w in words:
        src = w.get('source', 'IELTS General')
        if src not in source_groups:
            source_groups[src] = []
        source_groups[src].append(w)

    # Also group existing words (no source_file) by theme
    ungrouped = []
    for src, group in list(source_groups.items()):
        if src == 'IELTS General' or len(group) < 5:
            ungrouped.extend(group)
            del source_groups[src]

    # Build decks
    decks = []

    for src_name, src_words in source_groups.items():
        # Determine base title from source name
        base = os.path.splitext(src_name)[0]
        # Translate common file patterns to readable Chinese titles
        title_map = {
            '听力核心场景词汇': '听力场景词汇',
            '听力179词同义替换': '听力179同义替换考点词',
            '雅思阅读-538考点词表格整理': '阅读538考点词真经',
            '桃子99组雅思阅读同义词替换': '阅读99组同义词替换',
            '雅思口语练习题库': '口语练习核心词汇',
            '雅思听力地图题词汇汇总': '听力地图题词汇',
            '雅思词汇EXCEL版-顺序版': '雅思核心词汇（顺序版）',
            '雅思词汇EXCEL词-乱序版': '雅思核心词汇（乱序版）',
        }
        title = title_map.get(base, base)

        # Split into chunks of 24
        for i in range(0, len(src_words), DECK_SIZE):
            chunk = src_words[i:i + DECK_SIZE]
            if len(chunk) < 5:
                continue  # Skip tiny decks

            part = i // DECK_SIZE + 1
            total_parts = max(1, (len(src_words) + DECK_SIZE - 1) // DECK_SIZE)

            deck_id = slugify(src_name) + '-part-' + str(part)
            deck_title = title
            if total_parts > 1:
                deck_title += f' Part {part}'

            decks.append({
                'id': deck_id,
                'title': deck_title,
                'categoryTitle': title,
                'tag': 'Source-based',
                'part': part,
                'totalParts': total_parts,
                'wordCount': len(chunk),
                'words': chunk,
            })

    # Handle ungrouped (theme-based grouping for existing words)
    if ungrouped:
        theme_groups = OrderedDict()
        for w in ungrouped:
            theme = w.get('theme', 'Academic & Research')
            if theme not in theme_groups:
                theme_groups[theme] = []
            theme_groups[theme].append(w)

        for theme, theme_words in theme_groups.items():
            for i in range(0, len(theme_words), DECK_SIZE):
                chunk = theme_words[i:i + DECK_SIZE]
                if len(chunk) < 3:
                    continue

                part = i // DECK_SIZE + 1
                total_parts = max(1, (len(theme_words) + DECK_SIZE - 1) // DECK_SIZE)

                deck_id = slugify(theme) + '-part-' + str(part)
                deck_title = theme
                if total_parts > 1:
                    deck_title += f' Part {part}'

                decks.append({
                    'id': deck_id,
                    'title': deck_title,
                    'categoryTitle': theme,
                    'tag': theme,
                    'part': part,
                    'totalParts': total_parts,
                    'wordCount': len(chunk),
                    'words': chunk,
                })

    return decks


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    dry_run = '--dry-run' in sys.argv
    max_xls = int(next((a.split('=')[1] for a in sys.argv if a.startswith('--max-words=')), '5000'))

    print("=" * 70)
    print("IELTS Vocab Extraction & Deck Builder Pipeline")
    print(f"Started: {datetime.now().isoformat()}")
    print("=" * 70)

    # ── Load existing vocabulary ──
    existing_words = []
    if os.path.exists(VOCAB_JSON):
        with open(VOCAB_JSON, 'r', encoding='utf-8') as f:
            existing_words = json.load(f)
        print(f"\n📦 Existing vocabulary: {len(existing_words)} entries")

    # ── Phase 1: XLSX extraction ──
    print("\n" + "─" * 50)
    print("PHASE 1: XLSX Extraction")
    print("─" * 50)

    all_new_words = []

    # 1a: 听力核心场景词汇
    path = SOURCE_FILES.get('listen_scene', '')
    if os.path.exists(path):
        words = extract_listen_scene(path)
        all_new_words.extend(words)
    else:
        print(f"  ⚠ File not found: listen_scene")

    # 1b: 听力179词同义替换
    path = SOURCE_FILES.get('179_syn', '')
    if os.path.exists(path):
        words = extract_179_syn(path)
        all_new_words.extend(words)
    else:
        print(f"  ⚠ File not found: 179_syn")

    # 1c: 阅读538考点词
    path = SOURCE_FILES.get('538_read', '')
    if os.path.exists(path):
        words = extract_538_read(path)
        all_new_words.extend(words)
    else:
        print(f"  ⚠ File not found: 538_read")

    # ── Phase 2: DOCX extraction ──
    print("\n" + "─" * 50)
    print("PHASE 2: DOCX Extraction")
    print("─" * 50)

    for key in ['taozi_syn', 'speaking', 'map_vocab']:
        path = SOURCE_FILES.get(key, '')
        if os.path.exists(path):
            fname = os.path.basename(path)
            words = extract_docx_table(path, fname)
            all_new_words.extend(words)
        else:
            print(f"  ⚠ File not found: {key}")

    # ── Phase 3: PDF extraction (select key PDFs) ──
    print("\n" + "─" * 50)
    print("PHASE 3: PDF Extraction (text-based only)")
    print("─" * 50)

    pdf_targets = [
        ('awl', os.path.join(IELTS_LIBRARY, '02_词汇书与词表', '570 个英语学术核心词汇（AWL）.pdf')),
        ('listen_heqiong', os.path.join(IELTS_LIBRARY, '02_词汇书与词表', '何琼雅思听力必考词汇.pdf')),
        ('listen_classify', os.path.join(IELTS_LIBRARY, '02_词汇书与词表', '赠送-雅思听力分类词汇.pdf')),
        ('reading_classify', os.path.join(IELTS_LIBRARY, '02_词汇书与词表', '赠送-雅思阅读分类词汇.pdf.pdf')),
    ]

    for pdf_key, pdf_path in pdf_targets:
        if os.path.exists(pdf_path):
            fname = os.path.basename(pdf_path)
            words = extract_pdf_text(pdf_path, fname, max_pages=20)
            all_new_words.extend(words)
        else:
            print(f"  ⚠ File not found: {pdf_key}")

    # ── Phase 4: XLS extraction ──
    print("\n" + "─" * 50)
    print("PHASE 4: XLS Extraction")
    print("─" * 50)

    xls_targets = [
        ('xls_order', SOURCE_FILES.get('xls_order', '')),
        ('xls_random', SOURCE_FILES.get('xls_random', '')),
    ]

    # Check if xls_order and xls_random are the same content (just shuffled)
    # Only extract one if identical size
    xls_files_to_process = []
    for key, path in xls_targets:
        if path and os.path.exists(path):
            xls_files_to_process.append((key, path))

    # Only extract first XLS if both exist (to save time and avoid duplicates)
    if len(xls_files_to_process) >= 1:
        key, path = xls_files_to_process[0]
        fname = os.path.basename(path)
        words = extract_xls_vocab(path, fname, max_rows=max_xls)
        all_new_words.extend(words)

    # ── Normalize and filter ──
    print("\n" + "─" * 50)
    print("PHASE 5: Normalize, Filter & Dedup")
    print("─" * 50)

    raw_count = len(all_new_words)
    all_new_words = [normalize_word_entry(w) for w in all_new_words]
    all_new_words = [w for w in all_new_words if is_valid_entry(w)]
    print(f"  Raw extracted: {raw_count}")
    print(f"  After filtering: {len(all_new_words)}")

    # ── Merge with existing ──
    merged, added, skipped = merge_and_dedup(all_new_words, existing_words)
    print(f"  Existing words: {len(existing_words)}")
    print(f"  New words added: {added}")
    print(f"  Duplicates skipped: {skipped}")
    print(f"  TOTAL merged: {len(merged)}")

    # ── Build decks ──
    print("\n" + "─" * 50)
    print("PHASE 6: Build Decks")
    print("─" * 50)

    decks = build_decks(merged)
    total_deck_words = sum(d['wordCount'] for d in decks)
    print(f"  Total decks: {len(decks)}")
    print(f"  Total words in decks: {total_deck_words}")

    for d in decks:
        print(f"  📚 {d['title']}: {d['wordCount']} words (source: {d.get('categoryTitle', 'N/A')})")

    # ── Write output ──
    if not dry_run:
        print("\n" + "─" * 50)
        print("PHASE 7: Write Output")
        print("─" * 50)

        # Backup existing
        backup_path = VOCAB_JSON + '.backup.' + datetime.now().strftime('%Y%m%d_%H%M%S')
        if os.path.exists(VOCAB_JSON):
            import shutil
            shutil.copy2(VOCAB_JSON, backup_path)
            print(f"  📦 Backup created: {backup_path}")

        # Write merged vocabulary
        with open(VOCAB_JSON, 'w', encoding='utf-8') as f:
            json.dump(merged, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Written: {VOCAB_JSON} ({len(merged)} entries)")

        # Write decks JSON (for reference)
        decks_json_path = os.path.join(DATA_DIR, 'vocab_decks.json')
        decks_output = []
        for d in decks:
            # Don't include full word objects in decks file (too large)
            decks_output.append({
                'id': d['id'],
                'title': d['title'],
                'categoryTitle': d['categoryTitle'],
                'tag': d['tag'],
                'part': d['part'],
                'totalParts': d['totalParts'],
                'wordCount': d['wordCount'],
                'wordIds': [w['id'] for w in d['words']],
            })
        with open(decks_json_path, 'w', encoding='utf-8') as f:
            json.dump(decks_output, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Written: {decks_json_path} ({len(decks_output)} decks)")

        # Write summary
        summary = {
            'generated_at': datetime.now().isoformat(),
            'total_words': len(merged),
            'new_words_added': added,
            'duplicates_skipped': skipped,
            'total_decks': len(decks),
            'total_deck_words': total_deck_words,
            'decks': decks_output,
        }
        summary_path = os.path.join(DATA_DIR, 'vocab_extraction_summary.json')
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        print(f"  ✅ Written: {summary_path}")
    else:
        print("\n  🔍 DRY RUN — no files written")

    # ── Final report ──
    print("\n" + "=" * 70)
    print("EXTRACTION COMPLETE")
    print("=" * 70)
    print(f"  Total vocabulary entries: {len(merged)}")
    print(f"  New words added this run: {added}")
    print(f"  Total decks generated: {len(decks)}")
    print(f"  Words in decks: {total_deck_words}")
    print("=" * 70)

    return {
        'total_words': len(merged),
        'new_added': added,
        'total_decks': len(decks),
        'deck_words': total_deck_words,
    }


if __name__ == '__main__':
    main()
